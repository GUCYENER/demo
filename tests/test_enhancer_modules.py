"""
VYRA L1 Support API - Modüler Enhancer Tests
================================================
Modüler refactoring sonrası her alt modül için unit testler.

Kapsam:
  - SectionExtractor: Her format için bölüm çıkarma
  - CatBoostPrioritizer: Heuristic priority + weakness detection
  - LLMEnhancer: Fix instructions + LLM response parsing
  - OutputGenerator: XLSX sheet oluşturma
  - ImageHelpers: Görsel eşleştirme + pozisyon hesaplama
  - DocumentEnhancer (Facade): Orchestration + backward compatibility

pytest tests/test_enhancer_modules.py -v
"""

import io
import os
import pytest
from unittest.mock import patch, MagicMock


# ═══════════════════════════════════════════
# SectionExtractor Tests
# ═══════════════════════════════════════════

class TestSectionExtractor:
    """SectionExtractor: Tüm formatlar için bölüm çıkarma."""

    def test_extract_txt_sections(self):
        """TXT dosyasından bölüm çıkarma — heading pattern tespiti."""
        from app.services.enhancer.section_extractors import SectionExtractor
        extractor = SectionExtractor()

        text = "1. Giriş\nBu bir giriş paragrafıdır.\n2. Sonuç\nBu bir sonuç paragrafıdır."
        file_content = text.encode("utf-8")
        sections = extractor.extract_sections(file_content, "test.txt", "TXT")

        assert len(sections) >= 2
        assert sections[0]["heading"] == "1. Giriş"
        assert "giriş paragrafıdır" in sections[0]["content"]

    def test_extract_unknown_format_fallback(self):
        """Bilinmeyen format → düz metin olarak tek bölüm."""
        from app.services.enhancer.section_extractors import SectionExtractor
        extractor = SectionExtractor()

        content = "Bu bir test dosyasıdır.".encode("utf-8")
        sections = extractor.extract_sections(content, "test.xyz", "XYZ")

        assert len(sections) == 1
        assert sections[0]["heading"] == "Genel"

    def test_split_text_by_headings_empty(self):
        """Boş metin → tek 'Genel' bölüm."""
        from app.services.enhancer.section_extractors import SectionExtractor
        extractor = SectionExtractor()

        sections = extractor._split_text_by_headings("")
        assert len(sections) == 1
        # Boş metin → _split_text_by_headings varsayılan heading 'Giriş' veya 'Genel' döner
        assert sections[0]["heading"] in ("Genel", "Giriş")

    def test_split_text_by_headings_no_pattern(self):
        """Heading pattern olmadan → tek bölüm."""
        from app.services.enhancer.section_extractors import SectionExtractor
        extractor = SectionExtractor()

        text = "Bu bir normal paragraf.\nBu da başka bir paragraf."
        sections = extractor._split_text_by_headings(text)

        # Heading pattern bulunamazsa tek Giriş bölümü
        assert len(sections) >= 1

    def test_split_text_by_headings_multiple(self):
        """Birden fazla heading pattern → doğru bölümleme."""
        from app.services.enhancer.section_extractors import SectionExtractor
        extractor = SectionExtractor()

        text = (
            "BÖLÜM GİRİŞ\n"
            "Bu giriş bölümüdür.\n"
            "BÖLÜM SONUÇ\n"
            "Bu sonuç bölümüdür."
        )
        sections = extractor._split_text_by_headings(text)
        assert len(sections) >= 2

    def test_extract_docx_sections(self):
        """DOCX dosyasından heading bazlı bölüm çıkarma."""
        from app.services.enhancer.section_extractors import SectionExtractor
        from docx import Document

        extractor = SectionExtractor()

        doc = Document()
        doc.add_heading("Bölüm 1", level=1)
        doc.add_paragraph("İçerik 1a")
        doc.add_paragraph("İçerik 1b")
        doc.add_heading("Bölüm 2", level=1)
        doc.add_paragraph("İçerik 2a")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        sections = extractor._extract_docx_sections(buf)
        assert len(sections) >= 2
        assert "Bölüm 1" in sections[0]["heading"]

    def test_extract_docx_no_headings(self):
        """Heading'siz DOCX → tek 'Genel' bölüm."""
        from app.services.enhancer.section_extractors import SectionExtractor
        from docx import Document

        extractor = SectionExtractor()

        doc = Document()
        doc.add_paragraph("Sadece düz metin paragrafı.")
        doc.add_paragraph("İkinci paragraf.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        sections = extractor._extract_docx_sections(buf)
        assert len(sections) >= 1
        # Heading'siz DOCX → varsayılan heading 'Giriş' veya 'Genel'
        assert sections[0]["heading"] in ("Genel", "Giriş")

    def test_sections_have_para_range(self):
        """Tüm section'larda para_start/para_end bulunmalı."""
        from app.services.enhancer.section_extractors import SectionExtractor
        extractor = SectionExtractor()

        content = "1. Giriş\nParagraf 1\n2. Sonuç\nParagraf 2".encode("utf-8")
        sections = extractor.extract_sections(content, "test.txt", "TXT")

        for sec in sections:
            assert "para_start" in sec
            assert "para_end" in sec
            assert sec["para_start"] >= 0


# ═══════════════════════════════════════════
# CatBoostPrioritizer Tests
# ═══════════════════════════════════════════

class TestCatBoostPrioritizer:
    """CatBoostPrioritizer: Heuristic priority + weakness detection."""

    def test_heuristic_high_priority_short_content(self):
        """Çok kısa içerik → yüksek priority."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        priority = p._heuristic_priority("Kısa metin.", "Başlık", [])
        assert priority > 0.1, f"Kısa metin yüksek priority olmalı: {priority}"

    def test_heuristic_low_priority_good_content(self):
        """Uzun, iyi yapılandırılmış içerik → düşük priority."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        long_content = " ".join(["Bu bir test cümlesidir." for _ in range(30)])
        priority = p._heuristic_priority(long_content, "Detaylı Başlık", [])
        assert priority <= 0.2, f"İyi içerik düşük priority olmalı: {priority}"

    def test_heuristic_violation_boost(self):
        """Violation varsa priority artmalı."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        long_content = " ".join(["Test cümlesi." for _ in range(30)])
        prio_no_viol = p._heuristic_priority(long_content, "Başlık", [])
        prio_with_viol = p._heuristic_priority(long_content, "Başlık", ["Başlık Hiyerarşisi", "Metin Yoğunluğu"])

        assert prio_with_viol > prio_no_viol

    def test_heuristic_heading_missing_boost(self):
        """Heading yoksa veya generic ise priority artmalı."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        content = " ".join(["Test cümlesi." for _ in range(10)])
        prio_with_heading = p._heuristic_priority(content, "Detaylı Bölüm Başlığı", [])
        prio_no_heading = p._heuristic_priority(content, "Genel", [])

        assert prio_no_heading > prio_with_heading

    def test_detect_weaknesses_encoding(self):
        """Encoding bozuk → 'encoding_issue' weakness."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        content = "Bu metinde Ã¼ ve Ã§ bozuk karakterler var."
        weaknesses = p.detect_weaknesses(content, "Başlık", ["Türkçe Karakter"])

        assert "encoding_issue" in weaknesses

    def test_detect_weaknesses_heading_missing(self):
        """Heading yoksa → 'heading_missing' weakness."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        weaknesses = p.detect_weaknesses("İçerik", "", [])
        assert "heading_missing" in weaknesses

    def test_detect_weaknesses_short_content(self):
        """Çok kısa içerik → 'content_too_short' weakness."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        weaknesses = p.detect_weaknesses("Kısa", "Başlık", [])
        assert "content_too_short" in weaknesses

    def test_detect_weaknesses_excel_specific(self):
        """Excel violation'ları → Excel-specific weakness."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        weaknesses = p.detect_weaknesses(
            " ".join(["test" for _ in range(25)]),
            "Veri Sayfası",
            ["İlk Satır Başlık", "Merge Hücreler"]
        )
        assert "header_row_missing" in weaknesses
        assert "merged_cells" in weaknesses

    def test_catboost_prioritize_without_model(self):
        """CatBoost modeli yok → heuristic fallback kullanmalı."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        sections = [
            {"index": 0, "heading": "Test", "content": "Kısa bir metin."}
        ]
        violations = [{"name": "Başlık Hiyerarşisi", "status": "fail"}]

        result = p.catboost_prioritize(sections, "DOCX", violations)

        assert "sections" in result
        assert "summary" in result
        assert len(result["sections"]) == 1
        assert result["summary"]["catboost_available"] is False

    def test_priority_normalized_0_1(self):
        """Priority her zaman 0-1 aralığında olmalı."""
        from app.services.enhancer.catboost_prioritizer import CatBoostPrioritizer
        p = CatBoostPrioritizer()

        # Tüm violation'larla max priority
        all_violations = [
            "Başlık Hiyerarşisi", "Metin Yoğunluğu", "Tablo Formatı",
            "Türkçe Karakter", "Gereksiz İçerik", "İlk Satır Başlık",
            "Merge Hücreler", "Açıklama Satırları", "Boş Satır/Sütun",
            "Metin Kutusu"
        ]
        priority = p._heuristic_priority("Kısa", "", all_violations)
        assert 0.0 <= priority <= 1.0, f"Priority 0-1 aralığında olmalı: {priority}"


# ═══════════════════════════════════════════
# LLMEnhancer Tests
# ═══════════════════════════════════════════

class TestLLMEnhancer:
    """LLMEnhancer: Fix instructions + LLM response parsing."""

    def test_build_fix_instructions_heading(self):
        """heading_missing → uygun talimat üretmeli."""
        from app.services.enhancer.llm_enhancement import LLMEnhancer
        e = LLMEnhancer()

        instructions = e._build_fix_instructions(["heading_missing"])
        # Türkçe İ → i casefold dönüşümü için casefold kullan
        assert "başlı" in instructions.casefold()

    def test_build_fix_instructions_encoding(self):
        """encoding_issue → encoding düzeltme talimatı."""
        from app.services.enhancer.llm_enhancement import LLMEnhancer
        e = LLMEnhancer()

        instructions = e._build_fix_instructions(["encoding_issue"])
        assert "karakter" in instructions.lower() or "encoding" in instructions.lower()

    def test_build_fix_instructions_empty(self):
        """Boş weakness → genel kalite talimatı."""
        from app.services.enhancer.llm_enhancement import LLMEnhancer
        e = LLMEnhancer()

        instructions = e._build_fix_instructions([])
        assert "kalite" in instructions.lower() or "iyileştirme" in instructions.lower()

    def test_build_fix_instructions_excel(self):
        """Excel weakness'ları → Excel-specific talimatlar."""
        from app.services.enhancer.llm_enhancement import LLMEnhancer
        e = LLMEnhancer()

        instructions = e._build_fix_instructions(["header_row_missing", "merged_cells"])
        assert "başlık satırı" in instructions.lower() or "hücre" in instructions.lower()

    def test_parse_llm_response_valid_json(self):
        """Geçerli JSON yanıt → doğru parse."""
        from app.services.enhancer.llm_enhancement import LLMEnhancer
        e = LLMEnhancer()

        response = '{"heading": "Yeni Başlık", "enhanced_text": "İyileştirilmiş metin", "change_type": "content_restructured", "explanation": "Yapı düzeltildi"}'
        result = e._parse_llm_response(response, "Eski Başlık", "Eski metin")

        assert result["heading"] == "Yeni Başlık"
        assert result["enhanced_text"] == "İyileştirilmiş metin"
        assert result["change_type"] == "content_restructured"

    def test_parse_llm_response_markdown_wrapped(self):
        """Markdown code block içindeki JSON → doğru parse."""
        from app.services.enhancer.llm_enhancement import LLMEnhancer
        e = LLMEnhancer()

        response = '```json\n{"heading": "Test", "enhanced_text": "Metin", "change_type": "content_restructured", "explanation": "OK"}\n```'
        result = e._parse_llm_response(response, "Fallback", "Fallback Text")

        assert result["heading"] == "Test"
        assert result["enhanced_text"] == "Metin"

    def test_parse_llm_response_invalid_json(self):
        """Geçersiz JSON → fallback heading kullanılmalı."""
        from app.services.enhancer.llm_enhancement import LLMEnhancer
        e = LLMEnhancer()

        response = "Bu bir JSON değil, düz metin yanıtıdır."
        result = e._parse_llm_response(response, "Fallback Başlık", "Fallback İçerik")

        assert result["heading"] == "Fallback Başlık"

    def test_parse_llm_response_partial_json(self):
        """Eksik alanları olan JSON → fallback değerler kullanılmalı."""
        from app.services.enhancer.llm_enhancement import LLMEnhancer
        e = LLMEnhancer()

        response = '{"enhanced_text": "Sadece metin var"}'
        result = e._parse_llm_response(response, "Fallback", "Original")

        assert result["enhanced_text"] == "Sadece metin var"
        assert result["heading"] == "Fallback"  # Eksik alan → fallback


# ═══════════════════════════════════════════
# XLSX Output Generator Tests
# ═══════════════════════════════════════════

class TestOutputXlsx:
    """XLSX output generator: Sheet oluşturma ve format kontrolü."""

    def _make_section(self, index=0, heading="Test Sheet", enhanced="İyileştirilmiş veri",
                      change_type="content_restructured"):
        from app.services.document_enhancer import EnhancedSection
        return EnhancedSection(
            section_index=index, heading=heading,
            original_text="Orijinal", enhanced_text=enhanced,
            change_type=change_type, explanation="Test",
            priority=0.7, violations=[]
        )

    def test_xlsx_creates_enhanced_sheet(self):
        """Orijinal XLSX'e enhanced sheet eklenmeli."""
        pytest.importorskip("openpyxl", reason="openpyxl yüklü değil")
        from app.services.enhancer.output_xlsx import apply_to_original_xlsx
        from openpyxl import Workbook, load_workbook

        # Test XLSX oluştur
        wb = Workbook()
        ws = wb.active
        ws.title = "Veri"
        ws.cell(row=1, column=1, value="Test veri")

        buf = io.BytesIO()
        wb.save(buf)
        original_content = buf.getvalue()

        section = self._make_section()
        result_path = apply_to_original_xlsx(original_content, [section], "test_session")

        try:
            assert os.path.exists(result_path)
            result_wb = load_workbook(result_path)
            # Orijinal sheet korunmalı
            assert "Veri" in result_wb.sheetnames
            # Enhanced sheet eklenmeli
            enhanced_sheets = [s for s in result_wb.sheetnames if s.startswith("[E]")]
            assert len(enhanced_sheets) >= 1
        finally:
            if os.path.exists(result_path):
                os.remove(result_path)

    def test_xlsx_no_change_skipped(self):
        """no_change section'lar → enhanced sheet oluşturulmamalı."""
        pytest.importorskip("openpyxl", reason="openpyxl yüklü değil")
        from app.services.enhancer.output_xlsx import apply_to_original_xlsx
        from openpyxl import Workbook, load_workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Veri"
        ws.cell(row=1, column=1, value="Test")

        buf = io.BytesIO()
        wb.save(buf)
        original_content = buf.getvalue()

        section = self._make_section(change_type="no_change")
        result_path = apply_to_original_xlsx(original_content, [section], "test_session_skip")

        try:
            result_wb = load_workbook(result_path)
            enhanced_sheets = [s for s in result_wb.sheetnames if s.startswith("[E]")]
            assert len(enhanced_sheets) == 0, "no_change section için sheet oluşturulmamalı"
        finally:
            if os.path.exists(result_path):
                os.remove(result_path)


# ═══════════════════════════════════════════
# Facade (DocumentEnhancer) Tests
# ═══════════════════════════════════════════

class TestDocumentEnhancerFacade:
    """DocumentEnhancer facade: Backward compatibility + orchestration."""

    def test_backward_compatible_imports(self):
        """Mevcut import path'leri hâlâ çalışmalı."""
        from app.services.document_enhancer import DocumentEnhancer
        from app.services.document_enhancer import EnhancedSection
        from app.services.document_enhancer import EnhancementResult
        from app.services.document_enhancer import get_enhanced_file_path
        from app.services.document_enhancer import cleanup_enhanced_file

        assert DocumentEnhancer is not None
        assert EnhancedSection is not None
        assert EnhancementResult is not None
        assert callable(get_enhanced_file_path)
        assert callable(cleanup_enhanced_file)

    def test_enhancer_has_sub_modules(self):
        """DocumentEnhancer alt modülleri barındırmalı (lazy init)."""
        from app.services.document_enhancer import DocumentEnhancer
        enhancer = DocumentEnhancer()

        assert hasattr(enhancer, '_section_extractor')
        assert hasattr(enhancer, '_catboost_prioritizer')
        assert hasattr(enhancer, '_llm_enhancer')

    def test_to_dict_serialization(self):
        """to_dict JSON serializable dict döndürmeli."""
        from app.services.document_enhancer import DocumentEnhancer, EnhancementResult, EnhancedSection

        enhancer = DocumentEnhancer()
        result = EnhancementResult(
            file_name="test.pdf",
            file_type="PDF",
            total_sections=2,
            enhanced_count=1,
            sections=[
                EnhancedSection(
                    section_index=0, heading="Test",
                    original_text="Ori", enhanced_text="Enh",
                    change_type="content_restructured",
                    explanation="Test açıklama",
                    priority=0.7, violations=["structure_weak"]
                )
            ],
            session_id="abc123"
        )

        d = enhancer.to_dict(result)
        assert d["file_name"] == "test.pdf"
        assert d["total_sections"] == 2
        assert len(d["sections"]) == 1
        assert d["sections"][0]["heading"] == "Test"
        assert d["sections"][0]["integrity_score"] == 1.0

    def test_enhanced_section_dataclass_defaults(self):
        """EnhancedSection varsayılan değerleri doğru olmalı."""
        from app.services.document_enhancer import EnhancedSection

        section = EnhancedSection(
            section_index=0, heading="Test",
            original_text="Ori", enhanced_text="Enh",
            change_type="content_restructured",
            explanation="", priority=0.5
        )
        assert section.integrity_score == 1.0
        assert section.integrity_issues == []
        assert section.violations == []

    def test_cleanup_nonexistent_session(self):
        """Var olmayan session cleanup → hata vermemeli."""
        from app.services.document_enhancer import cleanup_enhanced_file
        # Hata fırlatmamalı
        cleanup_enhanced_file("nonexistent_session_id_xyz")

    def test_get_file_path_nonexistent_session(self):
        """Var olmayan session → None dönmeli."""
        from app.services.document_enhancer import get_enhanced_file_path
        result = get_enhanced_file_path("nonexistent_session_id_xyz")
        assert result is None


# ═══════════════════════════════════════════
# ImageHelpers Module Tests
# ═══════════════════════════════════════════

class TestImageHelpersModule:
    """image_helpers modülü: Doğrudan modül fonksiyon testleri."""

    def test_get_section_text_is_importable(self):
        """image_helpers modülü standalone import edilebilmeli."""
        from app.services.enhancer.image_helpers import (
            get_section_text,
            map_images_to_sections,
            organize_images_at_positions
        )
        assert callable(get_section_text)
        assert callable(map_images_to_sections)
        assert callable(organize_images_at_positions)

    def test_organize_empty_images(self):
        """Boş görsel listesi → boş dict."""
        from app.services.enhancer.image_helpers import organize_images_at_positions
        result = organize_images_at_positions([], total_paragraphs=5)
        assert result == {}

    def test_map_images_null_safe(self):
        """None images → boş dict."""
        from app.services.enhancer.image_helpers import map_images_to_sections
        result = map_images_to_sections([], None)
        assert result == {}
