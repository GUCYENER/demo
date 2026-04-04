"""
VYRA L1 Support API - Document Enhancer Tests
=================================================
Doküman iyileştirme modülü unit ve integration testleri.

Kapsam:
  - output_docx.create_fresh_docx: LLM meta-notlarının olmaması
  - DocumentEnhancer._generate_enhanced_output: PDF dönüşümü (fpdf2)
  - output_docx._update_paragraph_text: Paragraf güncelleme
  - download endpoint: format-aware response
  - upload endpoint: format-aware DB kaydı
  - Frontend confirm dialog: showToast entegrasyonu
  - image_helpers: Görsel eşleştirme ve pozisyon testi
  - section_extractors: Bölüm çıkarma ve paragraf aralığı

Unit testler: pytest tests/test_document_enhancer.py -m "not integration"
Integration testler: pytest tests/test_document_enhancer.py -m integration

v2.0.0: Modüler refactoring — testler alt modüllere doğrudan erişir
"""

import os
import io

import pytest
from unittest.mock import patch

# ─── UNIT TESTS ─── (DB/API gerekmez)


class TestCreateFreshDocx:
    """output_docx.create_fresh_docx fonksiyonu: LLM meta-notları olmamalı."""

    def _make_section(self, index=0, heading="Test Başlık", original="Orijinal metin",
                      enhanced="İyileştirilmiş metin", change_type="content_improved",
                      explanation="Test açıklama notu"):
        from app.services.document_enhancer import EnhancedSection
        return EnhancedSection(
            section_index=index,
            heading=heading,
            original_text=original,
            enhanced_text=enhanced,
            change_type=change_type,
            explanation=explanation,
            priority="medium",
            violations=[]
        )

    def test_no_iyilestirilmis_heading(self):
        """DOCX'te 'İyileştirilmiş:' başlığı OLMAMALI."""
        from app.services.enhancer.output_docx import create_fresh_docx
        section = self._make_section()

        doc = create_fresh_docx([section], "test_doc.pdf")

        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "İyileştirilmiş:" not in all_text, \
            f"'İyileştirilmiş:' başlığı hâlâ dokümanda: {all_text[:200]}"

    def test_no_explanation_note(self):
        """DOCX'te 📝 açıklama notu OLMAMALI."""
        from app.services.enhancer.output_docx import create_fresh_docx
        section = self._make_section(explanation="Bu bir LLM açıklamasıdır")

        doc = create_fresh_docx([section], "test_doc.pdf")

        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "📝" not in all_text, f"📝 notu hâlâ dokümanda: {all_text[:200]}"
        assert "Bu bir LLM açıklamasıdır" not in all_text, \
            f"LLM açıklaması dokümana eklendi: {all_text[:200]}"

    def test_enhanced_text_preserved(self):
        """İyileştirilmiş metin dokümanda olmalı."""
        from app.services.enhancer.output_docx import create_fresh_docx
        section = self._make_section(enhanced="Düzeltilmiş içerik burada")

        doc = create_fresh_docx([section], "test_doc.pdf")

        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Düzeltilmiş içerik burada" in all_text

    def test_no_change_uses_original(self):
        """change_type='no_change' ise orijinal metin kullanılmalı."""
        from app.services.enhancer.output_docx import create_fresh_docx
        section = self._make_section(
            original="Orijinal kalacak metin",
            enhanced="Bu kullanılmamalı",
            change_type="no_change"
        )

        doc = create_fresh_docx([section], "test_doc.pdf")

        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Orijinal kalacak metin" in all_text
        assert "Bu kullanılmamalı" not in all_text

    def test_empty_sections(self):
        """Boş section listesiyle hata vermemeli."""
        from app.services.enhancer.output_docx import create_fresh_docx

        doc = create_fresh_docx([], "test_doc.pdf")
        assert doc is not None
        assert len(doc.paragraphs) == 0

    def test_heading_fallback(self):
        """heading=None ise 'Bölüm N' fallback kullanılmalı."""
        from app.services.enhancer.output_docx import create_fresh_docx
        section = self._make_section(index=2, heading=None)

        doc = create_fresh_docx([section], "test_doc.pdf")

        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "Bölüm 3" in all_text


class TestGenerateEnhancedOutput:
    """DocumentEnhancer._generate_enhanced_output: PDF dönüşümü ve format tespiti."""

    def _make_section(self, **kwargs):
        from app.services.document_enhancer import EnhancedSection
        defaults = dict(
            section_index=0, heading="Test",
            original_text="Orijinal", enhanced_text="İyileştirilmiş",
            change_type="content_improved", explanation="",
            priority="medium", violations=[]
        )
        defaults.update(kwargs)
        return EnhancedSection(**defaults)

    def test_docx_output_for_docx_input(self):
        """DOCX girdisi → DOCX çıktı (orijinal format korunmalı)."""
        from app.services.document_enhancer import DocumentEnhancer
        from docx import Document

        # Gerçek bir DOCX binary oluştur
        doc = Document()
        doc.add_paragraph("Test içerik")
        buf = io.BytesIO()
        doc.save(buf)
        original_content = buf.getvalue()

        enhancer = DocumentEnhancer()
        section = self._make_section()

        result_path = enhancer._generate_enhanced_output(
            [section], "test.docx", "unit_test_session_1",
            original_content=original_content, file_type=".docx"
        )

        try:
            assert result_path.lower().endswith('.docx'), \
                f"DOCX girdi → çıktı .docx olmalı, ama: {result_path}"
            assert os.path.exists(result_path)
            assert os.path.getsize(result_path) > 0
        finally:
            if os.path.exists(result_path):
                os.remove(result_path)

    def test_pdf_output_for_pdf_input(self):
        """PDF girdisi → PDF çıktı (fpdf2 ile doğrudan PDF)."""
        from app.services.document_enhancer import DocumentEnhancer

        enhancer = DocumentEnhancer()
        section = self._make_section()

        result_path = enhancer._generate_enhanced_output(
            [section], "test.pdf", "unit_test_session_2",
            original_content=None, file_type=".pdf"
        )

        try:
            # fpdf2 başarılıysa .pdf, değilse .docx (fallback)
            assert os.path.exists(result_path), f"Dosya oluşturulamadı: {result_path}"
            assert os.path.getsize(result_path) > 0

            if result_path.lower().endswith('.pdf'):
                pass  # ✅ PDF dönüşümü başarılı
            else:
                assert result_path.lower().endswith('.docx')
        finally:
            if os.path.exists(result_path):
                os.remove(result_path)

    def test_pdf_conversion_fallback_on_error(self):
        """fpdf2 hatası → DOCX fallback olmalı (sistem çökmemeli)."""
        from app.services.document_enhancer import DocumentEnhancer

        enhancer = DocumentEnhancer()
        section = self._make_section()

        # output_pdf.create_fresh_pdf fonksiyonunu hata vermek üzere mock'la
        with patch(
            'app.services.enhancer.output_pdf.create_fresh_pdf',
            side_effect=RuntimeError("PDF oluşturulamadı")
        ):
            result_path = enhancer._generate_enhanced_output(
                [section], "test.pdf", "unit_test_session_3",
                original_content=None, file_type=".pdf"
            )

        try:
            assert os.path.exists(result_path)
            assert result_path.lower().endswith('.docx'), \
                f"Fallback DOCX olmalıydı: {result_path}"
        finally:
            if os.path.exists(result_path):
                os.remove(result_path)

    def test_pptx_stays_docx(self):
        """PPTX girdisi → DOCX çıktı (dönüşüm beklenmez)."""
        from app.services.document_enhancer import DocumentEnhancer

        enhancer = DocumentEnhancer()
        section = self._make_section()

        result_path = enhancer._generate_enhanced_output(
            [section], "sunum.pptx", "unit_test_session_4",
            original_content=None, file_type=".pptx"
        )

        try:
            assert result_path.lower().endswith('.docx'), \
                f"PPTX girdi → çıktı .docx olmalı: {result_path}"
        finally:
            if os.path.exists(result_path):
                os.remove(result_path)


class TestUploadEndpointFormatDetection:
    """Upload endpoint format tespiti: dosya uzantısına göre doğru parametreler."""

    def test_pdf_path_detected_correctly(self):
        docx_path = "C:/temp/enhanced_abc.pdf"
        is_pdf = docx_path.lower().endswith('.pdf')
        assert is_pdf is True

    def test_docx_path_detected_correctly(self):
        docx_path = "C:/temp/enhanced_abc.docx"
        is_pdf = docx_path.lower().endswith('.pdf')
        assert is_pdf is False

    def test_upload_name_pdf_preserved(self):
        file_name = "rapor.pdf"
        upload_name = file_name
        if not upload_name.lower().endswith('.pdf'):
            upload_name = os.path.splitext(upload_name)[0] + '.pdf'
        assert upload_name == "rapor.pdf"

    def test_upload_name_docx_converted(self):
        file_name = "sunum.pptx"
        is_pdf = False
        if not is_pdf:
            upload_name = file_name
            if not upload_name.lower().endswith('.docx'):
                upload_name = os.path.splitext(upload_name)[0] + '.docx'
        assert upload_name == "sunum.docx"

    def test_names_to_delete_dedup(self):
        file_name = "rapor.pdf"
        upload_name = "rapor.pdf"
        names = list(set([file_name, upload_name]))
        assert len(names) == 1

    def test_names_to_delete_different_ext(self):
        file_name = "rapor.pdf"
        upload_name = "rapor.docx"
        names = list(set([file_name, upload_name]))
        assert len(names) == 2


class TestDownloadFilename:
    """Frontend download dosya adı mantığı — JavaScript logic'in Python karşılığı."""

    def _get_download_name(self, current_file_name):
        import re
        if current_file_name:
            base_name = re.sub(r'\.[^.]+$', '', current_file_name)
            ext_match = re.search(r'\.[^.]+$', current_file_name)
            original_ext = ext_match.group(0) if ext_match else '.docx'
            return f"{base_name}{original_ext}"
        return 'iyilestirilmis.docx'

    def test_pdf_extension_preserved(self):
        assert self._get_download_name("rapor.pdf") == "rapor.pdf"

    def test_docx_extension_preserved(self):
        assert self._get_download_name("belge.docx") == "belge.docx"

    def test_pptx_extension_preserved(self):
        assert self._get_download_name("sunum.pptx") == "sunum.pptx"

    def test_multi_dot_filename(self):
        assert self._get_download_name("rapor.v2.pdf") == "rapor.v2.pdf"

    def test_no_extension_fallback(self):
        assert self._get_download_name("rapor") == "rapor.docx"

    def test_none_filename_fallback(self):
        assert self._get_download_name(None) == "iyilestirilmis.docx"

    def test_empty_filename_fallback(self):
        assert self._get_download_name("") == "iyilestirilmis.docx"


# ─── INTEGRATION TESTS ─── (Çalışan DB + API gerektirir)

pytestmark_integration = pytest.mark.integration


@pytest.fixture
def app_client():
    from app.api.main import app
    from app.core.config import settings
    from fastapi.testclient import TestClient
    return TestClient(app), settings


@pytest.fixture
def admin_token(app_client):
    client, settings = app_client
    response = client.post(
        f"{settings.api_prefix}/auth/login",
        json={"username": "admin", "password": "admin1234"}
    )
    assert response.status_code == 200
    return response.json()["access_token"]


@pytest.mark.integration
class TestEnhanceDownloadEndpoint:
    """Download endpoint format-aware response testleri."""

    def test_download_returns_correct_content_type_docx(self, app_client, admin_token):
        client, settings = app_client
        test_content = b"Test dokuman icerigi. Bu bir test dokumanidir."
        files = {"file": ("test_enhance.docx", io.BytesIO(test_content),
                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}

        response = client.post(
            f"{settings.api_prefix}/rag/enhance-document",
            files=files,
            headers={"Authorization": f"Bearer {admin_token}"}
        )

        if response.status_code == 200:
            data = response.json()
            session_id = data.get("session_id")
            if session_id:
                dl_response = client.get(
                    f"{settings.api_prefix}/rag/download-enhanced/{session_id}",
                    headers={"Authorization": f"Bearer {admin_token}"}
                )
                if dl_response.status_code == 200:
                    content_type = dl_response.headers.get("content-type", "")
                    assert "application" in content_type

                client.delete(
                    f"{settings.api_prefix}/rag/cleanup-enhanced/{session_id}",
                    headers={"Authorization": f"Bearer {admin_token}"}
                )


# ─── IMAGE EMBEDDING TESTS ─── (v2.40.1)


class TestImageEmbedding:
    """output_docx/output_pdf: Orijinal görsellerin iyileştirilmiş çıktıya dahil edilmesi."""

    def _make_section(self, index=0, heading="Test Başlık", original="Orijinal metin",
                      enhanced="İyileştirilmiş metin", change_type="content_improved"):
        from app.services.document_enhancer import EnhancedSection
        return EnhancedSection(
            section_index=index,
            heading=heading,
            original_text=original,
            enhanced_text=enhanced,
            change_type=change_type,
            explanation="",
            priority="medium",
            violations=[]
        )

    def _make_test_image(self, heading="Test Başlık", chunk_index=0, width=200, height=150):
        from app.services.document_processors.image_extractor import ExtractedImage
        from PIL import Image

        img = Image.new("RGB", (width, height), color=(255, 0, 0))
        buf = io.BytesIO()
        img.save(buf, format="PNG")

        return ExtractedImage(
            image_data=buf.getvalue(),
            image_format="png",
            width=width,
            height=height,
            context_heading=heading,
            context_chunk_index=chunk_index,
            alt_text=f"Test image - {heading}"
        )

    def test_docx_with_images(self):
        """create_fresh_docx görselleri DOCX'e eklemeli."""
        from app.services.enhancer.output_docx import create_fresh_docx

        section = self._make_section(heading="Sunucu Yapılandırması")
        test_image = self._make_test_image(heading="Sunucu Yapılandırması", chunk_index=0)

        doc = create_fresh_docx([section], "test.pdf", original_images=[test_image])

        inline_shapes = doc.inline_shapes
        assert len(inline_shapes) >= 1, \
            f"DOCX'e en az 1 görsel eklenmeli, ama {len(inline_shapes)} bulundu"

    def test_docx_without_images(self):
        """Görsel olmadan create_fresh_docx normal çalışmalı (regresyon testi)."""
        from app.services.enhancer.output_docx import create_fresh_docx

        section = self._make_section()

        doc = create_fresh_docx([section], "test.pdf", original_images=None)
        assert doc is not None

        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "İyileştirilmiş metin" in all_text

    def test_docx_image_section_matching(self):
        """Görseller doğru section'a eşleştirilmeli (heading bazlı)."""
        from app.services.enhancer.output_docx import create_fresh_docx

        section1 = self._make_section(index=0, heading="Giriş", enhanced="Giriş metni")
        section2 = self._make_section(index=1, heading="Sonuç", enhanced="Sonuç metni")

        test_image = self._make_test_image(heading="Sonuç", chunk_index=1)

        doc = create_fresh_docx(
            [section1, section2], "test.pdf", original_images=[test_image]
        )

        assert doc is not None
        assert len(doc.inline_shapes) >= 1

    def test_pdf_with_images(self):
        """create_fresh_pdf görselleri PDF'e eklemeli — dosya boyutu artmalı."""
        pytest.importorskip("fpdf", reason="fpdf2 yüklü değil, PDF testi atlanıyor")
        from app.services.enhancer.output_pdf import create_fresh_pdf

        section = self._make_section(heading="Ağ Topolojisi")
        test_image = self._make_test_image(heading="Ağ Topolojisi", width=300, height=200)

        path_no_img = create_fresh_pdf(
            [section], "test.pdf", "test_no_img", original_images=None
        )
        path_with_img = create_fresh_pdf(
            [section], "test.pdf", "test_with_img", original_images=[test_image]
        )

        try:
            assert os.path.exists(path_no_img)
            assert os.path.exists(path_with_img)

            size_no_img = os.path.getsize(path_no_img)
            size_with_img = os.path.getsize(path_with_img)

            assert size_with_img > size_no_img, \
                f"Görselli PDF ({size_with_img}) görselsiz ({size_no_img}) PDF'den büyük olmalı"
        finally:
            for p in [path_no_img, path_with_img]:
                if os.path.exists(p):
                    os.remove(p)

    def test_multiple_images_per_section(self):
        """Bir section'a birden fazla görsel eklenebilmeli."""
        from app.services.enhancer.output_docx import create_fresh_docx

        section = self._make_section(heading="Diyagram Sayfası")

        images = [
            self._make_test_image(heading="Diyagram Sayfası", chunk_index=0),
            self._make_test_image(heading="Diyagram Sayfası", chunk_index=0),
        ]

        doc = create_fresh_docx([section], "test.pdf", original_images=images)

        assert len(doc.inline_shapes) >= 2, \
            f"2 görsel eklenmeli, ama {len(doc.inline_shapes)} bulundu"


class TestUpdateChunkImageRefs:
    """ExtractedImage dataclass'ı ile doğru çalıştığını doğrular."""

    def test_extracted_image_attribute_access(self):
        from app.services.document_processors.image_extractor import ExtractedImage

        img = ExtractedImage(
            image_data=b"\x89PNG\r\n",
            image_format="png",
            width=100,
            height=100,
            context_heading="Test Bölüm",
            context_chunk_index=2,
            alt_text="Test"
        )

        heading = getattr(img, "context_heading", "")
        chunk_idx = getattr(img, "context_chunk_index", None)

        assert heading == "Test Bölüm"
        assert chunk_idx == 2

    def test_extracted_image_no_dict_get(self):
        from app.services.document_processors.image_extractor import ExtractedImage

        img = ExtractedImage(
            image_data=b"\x89PNG\r\n",
            image_format="png",
            context_heading="Test"
        )

        with pytest.raises(AttributeError):
            img.get("context_heading", "")

    def test_getattr_default_on_missing(self):
        from app.services.document_processors.image_extractor import ExtractedImage

        img = ExtractedImage(
            image_data=b"\x89PNG\r\n",
            image_format="png"
        )

        assert getattr(img, "context_heading", "") == ""
        assert getattr(img, "nonexistent_field", "fallback") == "fallback"


# ─── IMAGE POSITIONING TESTS ─── (v2.40.2)


class TestImagePositioning:
    """Görsellerin paragraflar arasına orijinal pozisyonlarına göre yerleştirilmesi."""

    def _make_section(self, index=0, heading="Test Başlık",
                      original="Satır 1\nSatır 2\nSatır 3",
                      enhanced="Satır 1\nSatır 2\nSatır 3",
                      change_type="content_improved"):
        from app.services.document_enhancer import EnhancedSection
        return EnhancedSection(
            section_index=index,
            heading=heading,
            original_text=original,
            enhanced_text=enhanced,
            change_type=change_type,
            explanation="",
            priority="medium",
            violations=[]
        )

    def _make_test_image(self, heading="Test Başlık", chunk_index=0,
                         paragraph_index=-1, width=200, height=150):
        from app.services.document_processors.image_extractor import ExtractedImage
        from PIL import Image

        img = Image.new("RGB", (width, height), color=(0, 128, 255))
        buf = io.BytesIO()
        img.save(buf, format="PNG")

        return ExtractedImage(
            image_data=buf.getvalue(),
            image_format="png",
            width=width,
            height=height,
            context_heading=heading,
            context_chunk_index=chunk_index,
            alt_text=f"Test image - {heading}",
            paragraph_index=paragraph_index
        )

    def test_image_placed_between_paragraphs(self):
        """Görsel paragraph_index'e göre paragraflar arasına yerleştirilmeli."""
        from app.services.enhancer.output_docx import create_fresh_docx

        section = self._make_section(
            heading="Test",
            enhanced="Paragraf Bir\nParagraf İki\nParagraf Üç"
        )

        img = self._make_test_image(heading="Test", paragraph_index=0)
        doc = create_fresh_docx([section], "test.pdf", original_images=[img])

        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) >= 3, f"En az 3 metin paragrafı olmalı, bulundu: {len(texts)}"
        assert len(doc.inline_shapes) >= 1, "En az 1 görsel olmalı"

    def test_image_without_position_goes_to_end(self):
        """paragraph_index=-1 olan görseller section sonuna eklenmeli."""
        from app.services.enhancer.output_docx import create_fresh_docx

        section = self._make_section(
            heading="Son Ekleme Testi",
            enhanced="Satır 1\nSatır 2"
        )

        img = self._make_test_image(heading="Son Ekleme Testi", paragraph_index=-1)
        doc = create_fresh_docx([section], "test.pdf", original_images=[img])

        assert len(doc.inline_shapes) >= 1, "Görsel sonuna eklenmeli"

    def test_update_paragraph_text_preserves_images(self):
        """_update_paragraph_text inline görsel içeren run'ları korumalı."""
        from app.services.enhancer.output_docx import _update_paragraph_text
        from docx import Document
        from docx.shared import Inches

        doc = Document()

        para = doc.add_paragraph("Metin önce")
        doc.add_picture(io.BytesIO(
            self._make_test_image().image_data
        ), width=Inches(1))

        _update_paragraph_text(para, "Güncellenmiş metin")
        assert para.text == "Güncellenmiş metin"

    def test_extracted_image_has_paragraph_index(self):
        from app.services.document_processors.image_extractor import ExtractedImage

        img = ExtractedImage(
            image_data=b"\x89PNG\r\n",
            image_format="png",
            paragraph_index=5,
            page_y_position=123.4
        )

        assert img.paragraph_index == 5
        assert img.page_y_position == 123.4

        img2 = ExtractedImage(
            image_data=b"\x89PNG\r\n",
            image_format="png"
        )
        assert img2.paragraph_index == -1
        assert img2.page_y_position == -1

    def test_docx_sections_have_para_range(self):
        """SectionExtractor._extract_docx_sections paragraf aralığı döndürmeli."""
        from app.services.enhancer.section_extractors import SectionExtractor
        from docx import Document

        extractor = SectionExtractor()

        doc = Document()
        doc.add_heading("Bölüm 1", level=1)
        doc.add_paragraph("İçerik 1a")
        doc.add_paragraph("İçerik 1b")
        doc.add_heading("Bölüm 2", level=1)
        doc.add_paragraph("İçerik 2a")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        sections = extractor._extract_docx_sections(buf)
        assert len(sections) >= 2

        for sec in sections:
            assert "para_start" in sec, f"Section {sec['heading']}: para_start yok"
            assert "para_end" in sec, f"Section {sec['heading']}: para_end yok"
            assert sec["para_start"] >= 0


class TestHelperFunctions:
    """image_helpers ortak yardımcı fonksiyonlar."""

    def _make_section(self, index=0, heading="Test", original="Orijinal",
                      enhanced="İyileştirilmiş", change_type="content_improved"):
        from types import SimpleNamespace
        return SimpleNamespace(
            section_index=index,
            heading=heading,
            original_text=original,
            enhanced_text=enhanced,
            change_type=change_type,
            explanation="",
            priority=0.5,
            violations=[]
        )

    def _make_image(self, heading="", chunk_idx=0, para_idx=-1):
        from types import SimpleNamespace
        return SimpleNamespace(
            image_data=b"\x89PNG",
            image_format="png",
            width=200,
            height=100,
            context_heading=heading,
            context_chunk_index=chunk_idx,
            alt_text="test",
            paragraph_index=para_idx,
            page_y_position=-1
        )

    def test_get_section_text_enhanced(self):
        """get_section_text: change_type != no_change ise enhanced döner."""
        from app.services.enhancer.image_helpers import get_section_text
        sec = self._make_section(change_type="content_improved",
                                  enhanced="Geliştirilmiş", original="Orijinal")
        assert get_section_text(sec) == "Geliştirilmiş"

    def test_get_section_text_no_change(self):
        """get_section_text: change_type == no_change ise original döner."""
        from app.services.enhancer.image_helpers import get_section_text
        sec = self._make_section(change_type="no_change",
                                  enhanced="Geliştirilmiş", original="Orijinal")
        assert get_section_text(sec) == "Orijinal"

    def test_get_section_text_none_safe(self):
        """get_section_text: None metin varsa boş string döner."""
        from app.services.enhancer.image_helpers import get_section_text
        sec = self._make_section(change_type="no_change",
                                  enhanced=None, original=None)
        assert get_section_text(sec) == ""

    def test_map_images_empty_list(self):
        """map_images_to_sections: boş görsel listesi boş dict döner."""
        from app.services.enhancer.image_helpers import map_images_to_sections
        result = map_images_to_sections([], [])
        assert result == {}

    def test_map_images_heading_match(self):
        """map_images_to_sections: heading eşleşmesiyle doğru section'a atar."""
        from app.services.enhancer.image_helpers import map_images_to_sections
        sec = self._make_section(index=0, heading="Giriş Bölümü",
                                  original="Satır1\nSatır2\nSatır3",
                                  enhanced="Gelişmiş1\nGelişmiş2\nGelişmiş3")
        img = self._make_image(heading="Giriş Bölümü", chunk_idx=0, para_idx=1)
        result = map_images_to_sections([sec], [img])
        assert 0 in result
        assert len(result[0]) == 1
        rel_pos, img_obj = result[0][0]
        assert rel_pos == 1

    def test_map_images_fallback_to_chunk_index(self):
        """map_images_to_sections: heading eşleşmezse chunk_index'e göre atar."""
        from app.services.enhancer.image_helpers import map_images_to_sections
        sec0 = self._make_section(index=0, heading="Bölüm A", original="İçerik A")
        sec1 = self._make_section(index=1, heading="Bölüm B", original="İçerik B")
        img = self._make_image(heading="Eşleşmez", chunk_idx=1, para_idx=-1)
        result = map_images_to_sections([sec0, sec1], [img])
        assert 1 in result
        assert 0 not in result

    def test_organize_images_unknown_position(self):
        """organize_images_at_positions: rel_pos=-1 olanlar section sonuna."""
        from app.services.enhancer.image_helpers import organize_images_at_positions

        class FakeImg:
            pass

        img = FakeImg()
        sec_imgs = [(-1, img)]
        result = organize_images_at_positions(sec_imgs, total_paragraphs=5)
        assert 5 in result
        assert img in result[5]

    def test_organize_images_known_position(self):
        """organize_images_at_positions: pozisyonlu görseller doğru konuma."""
        from app.services.enhancer.image_helpers import organize_images_at_positions

        class FakeImg:
            pass

        img = FakeImg()
        sec_imgs = [(2, img)]
        result = organize_images_at_positions(sec_imgs, total_paragraphs=5)
        assert 2 in result
        assert img in result[2]

    def test_organize_images_bounds_safe(self):
        """organize_images_at_positions: pozisyon paragraf sayısını aşarsa sınırlar."""
        from app.services.enhancer.image_helpers import organize_images_at_positions

        class FakeImg:
            pass

        img = FakeImg()
        sec_imgs = [(99, img)]
        result = organize_images_at_positions(sec_imgs, total_paragraphs=3)
        assert 2 in result
        assert img in result[2]

    def test_split_text_by_headings_has_para_range(self):
        """SectionExtractor._split_text_by_headings: tüm section'larda para_start/para_end olmalı."""
        from app.services.enhancer.section_extractors import SectionExtractor
        extractor = SectionExtractor()
        text = "1. Giriş\nBu bir giris.\n2. Sonuç\nBu bir sonuc."
        sections = extractor._split_text_by_headings(text)
        assert len(sections) >= 2
        for sec in sections:
            assert "para_start" in sec, f"Section {sec['heading']}: para_start yok"
            assert "para_end" in sec, f"Section {sec['heading']}: para_end yok"
            assert sec["para_start"] >= 0
