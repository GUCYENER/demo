"""
VYRA - DB Sorgu Sonucu Export API
===================================
Veritabanı sorgu sonuçlarını Excel, Word ve PDF formatında dışa aktarır.

Desteklenen formatlar:
  - Excel (.xlsx)  — openpyxl
  - Word  (.docx)  — python-docx
  - PDF   (.pdf)   — fpdf2

Güvenlik:
  - Tüm endpoint'ler JWT auth gerektirir
  - Max 500 satır export sınırı
  - Hassas sütun maskeleme (safe_sql_executor'dan devralınır)

Version: 4.0.0
"""

from __future__ import annotations

import io
import json
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.api.routes.auth import get_current_user
from app.services.logging_service import log_system_event, log_warning

router = APIRouter(prefix="/api/db", tags=["db-export"])

MAX_EXPORT_ROWS = 500


# =============================================================================
# REQUEST MODEL
# =============================================================================

class DBExportRequest(BaseModel):
    """Export isteği."""
    columns: List[str] = Field(..., description="Sütun adları listesi")
    rows: List[Dict[str, Any]] = Field(..., description="Satır verisi — dict listesi")
    title: Optional[str] = Field("VYRA Sorgu Sonucu", description="Dosya başlığı")
    query: Optional[str] = Field(None, description="Orijinal kullanıcı sorusu")
    sql: Optional[str] = Field(None, description="Çalıştırılan SQL")
    include_narrative: bool = Field(False, description="Word/PDF'e LLM yorumu ekle")


# =============================================================================
# EXCEL EXPORT
# =============================================================================

@router.post("/export/excel")
async def export_excel(
    request: DBExportRequest,
    user=Depends(get_current_user),
):
    """
    DB sorgu sonucunu Excel (.xlsx) olarak indir.
    - Header satırı kalın + koyu mavi arka plan
    - Kolonlar otomatik genişlik
    - Satır sayısı max 500
    """
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise HTTPException(status_code=500, detail="openpyxl kütüphanesi yüklü değil.")

    rows = request.rows[:MAX_EXPORT_ROWS]
    columns = request.columns

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = _safe_sheet_name(request.title or "Sonuç")

    # Başlık satırı (varsa)
    meta_row = 1
    ws.cell(row=meta_row, column=1, value=request.title or "VYRA Sorgu Sonucu")
    ws.cell(row=meta_row, column=1).font = Font(bold=True, size=13)
    if request.query:
        ws.cell(row=meta_row + 1, column=1, value=f"Soru: {request.query}")
        ws.cell(row=meta_row + 1, column=1).font = Font(italic=True, color="666666")
    ws.cell(row=meta_row + 2, column=1,
            value=f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')} | {len(rows)} kayıt")
    ws.cell(row=meta_row + 2, column=1).font = Font(color="888888", size=9)

    data_start_row = meta_row + 4  # Boşluk bırak

    # Header
    header_fill = PatternFill("solid", fgColor="1E3A5F")
    header_font = Font(color="FFFFFF", bold=True)
    for col_idx, col_name in enumerate(columns, 1):
        cell = ws.cell(row=data_start_row, column=col_idx, value=col_name)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Veri satırları
    alt_fill = PatternFill("solid", fgColor="F0F4FA")
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, col_name in enumerate(columns, 1):
            val = row_data.get(col_name, "")
            if val is None:
                val = ""
            cell = ws.cell(row=data_start_row + row_idx, column=col_idx, value=val)
            if row_idx % 2 == 0:
                cell.fill = alt_fill

    # Otomatik kolon genişliği
    for col_idx, col_name in enumerate(columns, 1):
        col_letter = get_column_letter(col_idx)
        max_len = len(str(col_name))
        for row_data in rows[:50]:  # İlk 50 satırdan genişlik hesapla
            val_str = str(row_data.get(col_name, "") or "")
            max_len = max(max_len, len(val_str))
        ws.column_dimensions[col_letter].width = min(max_len + 4, 50)

    # Satır yükseklikleri
    ws.row_dimensions[data_start_row].height = 22

    # SQL notu (varsa) — son satıra
    if request.sql:
        last_row = data_start_row + len(rows) + 2
        ws.cell(row=last_row, column=1, value=f"SQL: {request.sql[:500]}")
        ws.cell(row=last_row, column=1).font = Font(color="AAAAAA", size=8, italic=True)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    filename = _safe_filename(request.title, "xlsx")
    log_system_event("INFO", f"Excel export: {len(rows)} satır | user={user['id']}", "db_export")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# =============================================================================
# WORD EXPORT
# =============================================================================

@router.post("/export/word")
async def export_word(
    request: DBExportRequest,
    user=Depends(get_current_user),
):
    """
    DB sorgu sonucunu Word (.docx) olarak indir.
    - Başlık, tarih, soru, tablo
    - include_narrative=True → LLM yorumu eklenir
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx kütüphanesi yüklü değil.")

    rows = request.rows[:MAX_EXPORT_ROWS]
    columns = request.columns

    doc = Document()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Başlık
    title_para = doc.add_heading(request.title or "VYRA Sorgu Sonucu", level=1)
    title_para.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # Meta bilgi
    meta = doc.add_paragraph()
    meta.add_run(f"📅 {datetime.now().strftime('%d.%m.%Y %H:%M')}   ").font.size = Pt(9)
    meta.add_run(f"📊 {len(rows)} kayıt").font.size = Pt(9)
    meta.paragraph_format.space_after = Pt(6)

    if request.query:
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Soru: {request.query}")
        q_run.italic = True
        q_run.font.size = Pt(10)
        q_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        q_para.paragraph_format.space_after = Pt(10)

    # LLM narrative (opsiyonel)
    if request.include_narrative and request.query and rows:
        narrative = _generate_export_narrative(request.query, rows, columns)
        if narrative:
            doc.add_heading("Değerlendirme", level=2)
            narr_para = doc.add_paragraph(narrative)
            narr_para.paragraph_format.space_after = Pt(12)

    # Veri tablosu
    doc.add_heading("Veri", level=2)
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.style = "Light List Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(columns):
        hdr_cells[i].text = col_name
        run = hdr_cells[i].paragraphs[0].runs
        if run:
            run[0].bold = True
            run[0].font.size = Pt(9)

    # Veri satırları
    for row_idx, row_data in enumerate(rows, 1):
        cells = table.rows[row_idx].cells
        for col_idx, col_name in enumerate(columns):
            val = row_data.get(col_name, "")
            cells[col_idx].text = str(val) if val is not None else ""
            cells[col_idx].paragraphs[0].runs[0].font.size = Pt(8) if cells[col_idx].paragraphs[0].runs else None

    # SQL notu
    if request.sql:
        doc.add_paragraph()
        sql_para = doc.add_paragraph()
        sql_run = sql_para.add_run(f"Çalıştırılan SQL: {request.sql[:500]}")
        sql_run.font.size = Pt(7)
        sql_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        sql_run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = _safe_filename(request.title, "docx")
    log_system_event("INFO", f"Word export: {len(rows)} satır | user={user['id']}", "db_export")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# =============================================================================
# PDF EXPORT
# =============================================================================

@router.post("/export/pdf")
async def export_pdf(
    request: DBExportRequest,
    user=Depends(get_current_user),
):
    """
    DB sorgu sonucunu PDF olarak indir.
    - Türkçe karakter desteği (fpdf2 unicode)
    - Başlık, meta, tablo
    - include_narrative=True → LLM yorumu eklenir
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(status_code=500, detail="fpdf2 kütüphanesi yüklü değil.")

    rows = request.rows[:MAX_EXPORT_ROWS]
    columns = request.columns

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # v3.13.1: Türkçe karakter desteği — Unicode TTF font kaydet
    font_family, _safe = _register_pdf_unicode_font(pdf)
    # Aşağıda Helvetica yerine font_family kullanılır. _safe(text) latin-1
    # fallback durumunda Türkçe karakterleri ASCII'ye çevirir.

    pdf.set_font(font_family, size=16, style="B")
    pdf.set_text_color(30, 58, 95)  # #1E3A5F
    pdf.cell(0, 10, _safe(request.title or "VYRA Sorgu Sonucu"), ln=True)

    pdf.set_font(font_family, size=9)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 6, _safe(f"{datetime.now().strftime('%d.%m.%Y %H:%M')}  |  {len(rows)} kayıt"), ln=True)

    if request.query:
        pdf.set_font(font_family, style="I", size=10)
        pdf.set_text_color(60, 60, 60)
        # Uzun sorguyu kes
        q_short = request.query[:200] + ("..." if len(request.query) > 200 else "")
        pdf.multi_cell(0, 7, _safe(f"Soru: {q_short}"))
    pdf.ln(4)

    # LLM narrative (opsiyonel)
    if request.include_narrative and request.query and rows:
        narrative = _generate_export_narrative(request.query, rows, columns)
        if narrative:
            pdf.set_font(font_family, style="B", size=11)
            pdf.set_text_color(30, 58, 95)
            pdf.cell(0, 8, _safe("Değerlendirme"), ln=True)
            pdf.set_font(font_family, size=9)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(0, 6, _safe(narrative[:1000]))
            pdf.ln(4)

    # Tablo
    pdf.set_font(font_family, style="B", size=11)
    pdf.set_text_color(30, 58, 95)
    pdf.cell(0, 8, _safe("Veri"), ln=True)
    pdf.ln(2)

    # Kolon genişliklerini hesapla
    page_width = pdf.w - 2 * pdf.l_margin
    col_count = len(columns)
    col_w = min(page_width / col_count, 55) if col_count > 0 else page_width
    row_h = 6

    # Header
    pdf.set_fill_color(30, 58, 95)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font(font_family, style="B", size=8)
    for col_name in columns:
        pdf.cell(col_w, row_h + 1, _safe(_trunc(col_name, 18)), border=1, fill=True, align="C")
    pdf.ln()

    # Veri satırları
    pdf.set_font(font_family, size=7)
    for row_idx, row_data in enumerate(rows):
        if row_idx % 2 == 0:
            pdf.set_fill_color(240, 244, 250)
        else:
            pdf.set_fill_color(255, 255, 255)
        pdf.set_text_color(40, 40, 40)
        for col_name in columns:
            val = str(row_data.get(col_name, "") or "")
            pdf.cell(col_w, row_h, _safe(_trunc(val, 20)), border=1, fill=True)
        pdf.ln()

    # SQL notu
    if request.sql:
        pdf.ln(4)
        pdf.set_font(font_family, style="I", size=7)
        pdf.set_text_color(170, 170, 170)
        pdf.multi_cell(0, 5, _safe(f"SQL: {request.sql[:300]}"))

    buffer = io.BytesIO(pdf.output())
    buffer.seek(0)

    filename = _safe_filename(request.title, "pdf")
    log_system_event("INFO", f"PDF export: {len(rows)} satır | user={user['id']}", "db_export")

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# =============================================================================
# YARDIMCI FONKSİYONLAR
# =============================================================================

def _safe_filename(title: Optional[str], ext: str) -> str:
    """Dosya adını güvenli hale getirir."""
    base = title or "vyra_export"
    base = re.sub(r'[^\w\s-]', '', base).strip().replace(' ', '_')
    base = base[:40] or "vyra_export"
    date_str = datetime.now().strftime("%Y%m%d_%H%M")
    return f"{base}_{date_str}.{ext}"


def _safe_sheet_name(name: str) -> str:
    """Excel sheet adından geçersiz karakterleri temizler (max 31 karakter)."""
    name = re.sub(r'[\\/:*?\[\]]', '', name)
    return name[:31] or "Sonuc"


def _trunc(text: str, max_len: int) -> str:
    """PDF hücreleri için metni kes."""
    if len(text) <= max_len:
        return text
    return text[:max_len - 1] + "…"


# v3.13.1: PDF Unicode font cache
_PDF_FONT_REGISTERED: Dict[str, bool] = {}

def _register_pdf_unicode_font(pdf):
    """
    PDF için Unicode TTF font kaydeder (Türkçe karakter desteği).
    Returns: (font_family_name, safe_text_fn)

    Strateji:
      1) Windows: C:/Windows/Fonts/arial.ttf (+bd, +i, +bi)
      2) Linux/diğer: DejaVuSans varsa
      3) Hiçbiri yoksa: Helvetica (Latin-1) + ASCII transliterate fallback
    """
    import os
    candidates = [
        # (family_name, regular, bold, italic, bold_italic)
        ("DejaVuTr", "C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf",
         "C:/Windows/Fonts/ariali.ttf", "C:/Windows/Fonts/arialbi.ttf"),
        ("DejaVuTr", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf"),
    ]
    for family, reg, bld, ital, bi in candidates:
        if reg and os.path.exists(reg):
            try:
                pdf.add_font(family, "", reg, uni=True)
                if bld and os.path.exists(bld):
                    pdf.add_font(family, "B", bld, uni=True)
                if ital and os.path.exists(ital):
                    pdf.add_font(family, "I", ital, uni=True)
                if bi and os.path.exists(bi):
                    pdf.add_font(family, "BI", bi, uni=True)
                return family, (lambda s: s if s is not None else "")
            except Exception as _e:
                log_warning(f"PDF font kayit hatasi ({family}): {_e}", "db_export")
                continue

    # Fallback: Helvetica + ASCII transliteration
    _tr_map = str.maketrans({
        "ı": "i", "İ": "I", "ş": "s", "Ş": "S",
        "ğ": "g", "Ğ": "G", "ü": "u", "Ü": "U",
        "ö": "o", "Ö": "O", "ç": "c", "Ç": "C",
        "â": "a", "Â": "A", "î": "i", "Î": "I", "û": "u", "Û": "U",
        "—": "-", "–": "-", "…": "...", "“": '"', "”": '"', "‘": "'", "’": "'",
    })
    def _safe(s):
        if s is None:
            return ""
        try:
            s2 = str(s).translate(_tr_map)
            # Latin-1 dışı her şeyi at
            return s2.encode("latin-1", "replace").decode("latin-1")
        except Exception:
            return ""
    return "Helvetica", _safe


def _generate_export_narrative(query: str, rows: list, columns: list) -> str:
    """
    LLM ile veri yorumu üretir (Word/PDF export için).
    Hata durumunda boş string döner.
    """
    try:
        from app.core.llm import call_llm_api
        col_sample = ", ".join(columns[:8])
        data_sample = json.dumps(rows[:10], ensure_ascii=False, default=str)
        messages = [
            {
                "role": "system",
                "content": (
                    "Sen bir veri analiz uzmanısın. Verilen sorgu sonucunu "
                    "kısa, profesyonel ve anlaşılır Türkçe bir yönetici özeti olarak yaz. "
                    "Max 3 paragraf. Bullet point kullanma. Sadece özet metin döndür."
                ),
            },
            {
                "role": "user",
                "content": (
                    f"Kullanıcı sorusu: {query}\n"
                    f"Sütunlar: {col_sample}\n"
                    f"Örnek veri (ilk 10 satır):\n{data_sample}"
                ),
            },
        ]
        result = call_llm_api(messages, temperature=0.3)
        return result or ""
    except Exception as e:
        log_warning(f"Export narrative üretim hatası: {e}", "db_export")
        return ""
