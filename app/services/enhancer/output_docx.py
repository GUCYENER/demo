"""
VYRA L1 Support API - DOCX Output Generator
==============================================
DOCX oluşturma ve düzenleme.
Orijinal DOCX güncelleme + sıfırdan DOCX oluşturma + selective download.

Author: VYRA AI Team
Version: 1.0.0 (v3.3.3)
"""

import io
from typing import Dict, List

from app.services.logging_service import log_warning
from app.services.enhancer.image_helpers import (
    get_section_text, map_images_to_sections, organize_images_at_positions
)

from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from app.services.document_enhancer import EnhancedSection


def apply_to_original_docx(original_content: bytes, sections: List['EnhancedSection']):
    """
    Orijinal DOCX'i açıp sadece iyileştirilmiş bölümlerin paragraflarını günceller.
    Stiller, tablolar, resimler, header/footer korunur.
    """
    from docx import Document

    doc = Document(io.BytesIO(original_content))

    # Değişiklik gereken section'ları filtrele
    changed_sections = {
        s.section_index: s for s in sections if s.change_type != "no_change"
    }

    if not changed_sections:
        return doc  # Değişiklik yok, orijinali aynen döndür

    # Paragrafları heading bazlı section'lara eşle
    current_heading = "Giriş"
    section_idx = 0
    section_paragraphs: Dict[int, List] = {}  # section_idx → [paragraph_refs]
    section_headings: Dict[int, str] = {}  # section_idx → heading_text

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading") and para.text.strip():
            # Önceki section'ı kapat
            if section_idx > 0 or section_paragraphs.get(0):
                section_idx += 1 if section_paragraphs.get(section_idx) else 0

            current_heading = para.text.strip()
            section_headings[section_idx] = current_heading

            # Heading paragrafı eğer iyileştirilmiş heading varsa güncelle
            if section_idx in changed_sections:
                new_heading = changed_sections[section_idx].heading
                if new_heading and new_heading != current_heading:
                    _update_paragraph_text(para, new_heading)

            # Bir sonraki section için hazırla
            section_idx += 1
            section_paragraphs.setdefault(section_idx, [])
        else:
            section_paragraphs.setdefault(section_idx, []).append(para)

    # Her değişen section için paragrafları güncelle
    for sec_idx, enhanced in changed_sections.items():
        paragraphs = section_paragraphs.get(sec_idx, [])
        if not paragraphs:
            continue

        # İyileştirilmiş metni satırlara böl
        new_lines = [line for line in enhanced.enhanced_text.split("\n") if line.strip()]

        # Mevcut paragrafları güncelle veya ekle
        for i, para in enumerate(paragraphs):
            if i < len(new_lines):
                # Paragraf var → metnini güncelle (stil korunur)
                _update_paragraph_text(para, new_lines[i])
            else:
                # Fazla paragrafları temizle (boşalt ama silme — stil korunsun)
                _update_paragraph_text(para, "")

        # Yeni metin paragraflardan fazlaysa kalan satırları son paragrafın
        # sonuna ekle (doküman yapısını bozmamak için)
        if len(new_lines) > len(paragraphs) and paragraphs:
            last_para = paragraphs[-1]
            remaining = new_lines[len(paragraphs):]
            current_text = last_para.text
            combined = current_text + "\n" + "\n".join(remaining) if current_text else "\n".join(remaining)
            _update_paragraph_text(last_para, combined)

    return doc


def _update_paragraph_text(para, new_text: str):
    """
    Paragrafın metnini güncellerken orijinal stil ve formatı korur.
    İlk metin run'ının formatını referans alarak metin run'larını günceller.

    KRİTİK: Inline görsel (drawing) içeren run'lar ASLA temizlenmez —
    bu sayede orijinal DOCX'teki görseller korunur.
    """
    if not para.runs:
        # Run yoksa direkt text güncelle
        para.text = new_text
        return

    # Run'ları iki gruba ayır: metin run'ları ve görsel (drawing) run'ları
    # OOXML namespace'leri: w:drawing ve wp:inline
    WML_DRAWING = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
    WP_INLINE = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'

    text_runs = []
    image_runs = []

    for run in para.runs:
        has_drawing = (
            run._element.findall(f'.//{WML_DRAWING}') or
            run._element.findall(f'.//{WP_INLINE}')
        )
        if has_drawing:
            image_runs.append(run)
        else:
            text_runs.append(run)

    # Eğer paragrafta HİÇ metin run'ı yoksa (paragraf sadece görsel ise)
    # → hiçbir şey yapma, görseli koru
    if not text_runs:
        return

    # İlk metin run'ından font özelliklerini sakla
    first_run = text_runs[0]
    font_props = {}
    for prop in ('bold', 'italic', 'underline', 'size', 'name'):
        val = getattr(first_run.font, prop, None)
        if val is not None:
            font_props[prop] = val

    # Color'u da sakla
    try:
        if first_run.font.color and first_run.font.color.rgb:
            font_props['color_rgb'] = first_run.font.color.rgb
    except Exception as e:
        log_warning(f"Font color okuma hatası: {e}", "enhancer")

    # Sadece METIN run'larını temizle (görsel run'larına DOKUNMA)
    for run in text_runs:
        run.clear()

    # İlk metin run'ının metnini güncelle
    first_run.text = new_text

    # Saklanan font özelliklerini geri yükle
    for prop, val in font_props.items():
        if prop == 'color_rgb':
            try:
                first_run.font.color.rgb = val
            except Exception as e:
                log_warning(f"Font color geri yükleme hatası: {e}", "enhancer")
        else:
            try:
                setattr(first_run.font, prop, val)
            except Exception as e:
                log_warning(f"Font özelliği geri yükleme hatası ({prop}): {e}", "enhancer")


def create_fresh_docx(
    sections: List['EnhancedSection'],
    original_name: str,
    original_images: list = None
):
    """
    PDF/PPTX/TXT gibi DOCX olmayan dosyalar için sıfırdan DOCX oluşturur.
    Orijinal yapıyı korur — LLM meta-notları veya açıklama eklenmez.

    v2.40.2: Görseller bölüm sonuna topluca değil, paragraflar arasına
    orijinal sırasına göre (paragraph_index bazlı) yerleştirilir.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Görselleri bölümlere ve bölüm içi pozisyonlara eşleştir
    section_image_map = map_images_to_sections(sections, original_images)

    # Dökümanı oluştur
    for section in sections:
        heading_text = section.heading or f"Bölüm {section.section_index + 1}"
        doc.add_heading(heading_text, level=1)

        text = get_section_text(section)
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
        sec_imgs = section_image_map.get(section.section_index, [])

        # Görselleri paragraf pozisyonlarına organize et
        imgs_at_pos = organize_images_at_positions(sec_imgs, len(paragraphs))

        # Paragrafları ve görselleri sırayla ekle
        for p_idx, paragraph_text in enumerate(paragraphs):
            para = doc.add_paragraph(paragraph_text)
            para.style.font.size = Pt(11)

            # Bu paragraftan sonra eklenecek görseller var mı?
            for img_obj in imgs_at_pos.get(p_idx, []):
                _add_image_to_docx(doc, img_obj)

        # Section sonuna eklenmesi gereken görseller (pozisyonsuz)
        for img_obj in imgs_at_pos.get(len(paragraphs), []):
            _add_image_to_docx(doc, img_obj)

    return doc


def _add_image_to_docx(doc, img_obj):
    """Tek bir görseli DOCX'e ekler — boyut ölçekleme ve hata koruması dahil."""
    from docx.shared import Inches
    try:
        img_data = getattr(img_obj, "image_data", None)
        if not img_data:
            return

        img_stream = io.BytesIO(img_data)

        # Boyut hesapla: max 5.5 inch genişlik (A4 margins dahil)
        width = getattr(img_obj, "width", 0) or 400
        max_width_inch = 5.5
        w_inch = min(width / 96.0, max_width_inch)

        doc.add_picture(img_stream, width=Inches(w_inch))
    except Exception as e:
        log_warning(f"DOCX görsel ekleme hatası: {e}", "enhancer")
