"""
DOCX Document Processor - v3.0
==============================
Semantik chunking ile geliştirilmiş DOCX işleme.
- Max chunk size kontrolü
- Recursive paragraph/sentence bölme
- Türkçe pattern tanıma
"""

import logging
import re
from pathlib import Path
from typing import BinaryIO, List, Dict, Any

from .base import BaseDocumentProcessor

logger = logging.getLogger("vyra")


class DOCXProcessor(BaseDocumentProcessor):
    """DOCX dosyalarını işleyen processor - Semantik Chunking v3.0"""
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    PROCESSOR_NAME = "DOCXProcessor"
    
    # Chunk boyut limitleri
    MAX_CHUNK_SIZE = 2000   # Karakter üst limiti
    MIN_CHUNK_SIZE = 100   # Çok küçük chunk'ları önle
    OVERLAP_SIZE = 50      # Overlap için
    
    # Türkçe doküman pattern'leri (semantik marker'lar)
    SEMANTIC_MARKERS = [
        r'^➤\s',           # Bullet with arrow
        r'^>\s',           # Quote/bullet
        r'^•\s',           # Bullet
        r'^-\s',           # Dash bullet
        r'^\d+\.\s',       # Numbered list (1. 2. 3.)
        r'^[a-zA-Z]\)\s',  # Letter list (a) b) c))
        r'^\(\d+\)\s',     # Parenthesized number ((1) (2))
        r'^NOT:\s',        # Note marker
        r'^ADRES:\s',      # Address marker
        r'^https?://',     # URL
        r'^[A-ZÇĞİÖŞÜ][A-ZÇĞİÖŞÜa-zçğıöşü\s]+:\s',  # Key: value pattern
    ]
    
    def extract_text(self, file_path: Path) -> str:
        """DOCX dosyasından metin çıkarır"""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            return self._extract_from_document(doc)
            
        except ImportError:
            raise ImportError("python-docx kütüphanesi yüklü değil. 'pip install python-docx' komutunu çalıştırın.")
        except Exception as e:
            raise RuntimeError(f"DOCX işleme hatası: {str(e)}")
    
    def extract_text_from_bytes(self, file_obj: BinaryIO, file_name: str) -> str:
        """BytesIO'dan DOCX metni çıkarır"""
        try:
            from docx import Document
            
            doc = Document(file_obj)
            return self._extract_from_document(doc)
            
        except ImportError:
            raise ImportError("python-docx kütüphanesi yüklü değil.")
        except Exception as e:
            raise RuntimeError(f"DOCX işleme hatası: {str(e)}")
    
    def extract_chunks(self, file_path: Path = None, file_obj: BinaryIO = None, file_name: str = None) -> List[Dict[str, Any]]:
        """
        DOCX dosyasından semantik chunk'lar çıkarır.
        - Heading bazlı bölme
        - Max size kontrolü ile recursive splitting
        - Tablolar satır bazlı chunk
        
        Returns:
            List of {"text": "...", "metadata": {...}}
        """
        try:
            from docx import Document
            
            if file_path:
                doc = Document(str(file_path))
            elif file_obj:
                doc = Document(file_obj)
            else:
                raise ValueError("file_path veya file_obj gerekli")
            
            return self._chunks_from_document(doc)
            
        except ImportError:
            raise ImportError("python-docx kütüphanesi yüklü değil.")
        except Exception as e:
            raise RuntimeError(f"DOCX işleme hatası: {str(e)}")
    
    def _split_large_chunk(self, text: str, max_size: int = None) -> List[str]:
        """
        Büyük chunk'ları anlamlı noktalardan böler.
        Öncelik sırası: paragraf > semantik marker > cümle sonu > karakter limiti
        
        Args:
            text: Bölünecek metin
            max_size: Maximum chunk boyutu (karakter)
            
        Returns:
            List of text chunks
        """
        max_size = max_size or self.MAX_CHUNK_SIZE
        
        # Zaten küçükse direkt döndür
        if len(text) <= max_size:
            return [text] if len(text.strip()) >= self.MIN_CHUNK_SIZE else []
        
        chunks = []
        
        # 1. Çift newline ile paragraf bazlı bölme dene
        paragraphs = text.split('\n\n')
        if len(paragraphs) > 1:
            current = ""
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                    
                test_len = len(current) + len(para) + (2 if current else 0)
                
                if test_len <= max_size:
                    current += ("\n\n" if current else "") + para
                else:
                    # Mevcut chunk'ı kaydet
                    if current:
                        sub_chunks = self._split_large_chunk(current, max_size)
                        chunks.extend(sub_chunks)
                    current = para
            
            # Son parçayı kaydet
            if current:
                sub_chunks = self._split_large_chunk(current, max_size)
                chunks.extend(sub_chunks)
            
            return chunks
        
        # 2. Tek newline ile satır bazlı bölme
        lines = text.split('\n')
        if len(lines) > 1:
            current = ""
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Semantik marker kontrolü - yeni bölüm başlangıcı mı?
                is_new_section = any(re.match(pattern, line) for pattern in self.SEMANTIC_MARKERS)
                
                test_len = len(current) + len(line) + (1 if current else 0)
                
                # Yeni semantik section veya size aşımı
                if (is_new_section and len(current) >= self.MIN_CHUNK_SIZE) or test_len > max_size:
                    if current:
                        sub_chunks = self._split_large_chunk(current, max_size)
                        chunks.extend(sub_chunks)
                    current = line
                else:
                    current += ("\n" if current else "") + line
            
            if current:
                sub_chunks = self._split_large_chunk(current, max_size)
                chunks.extend(sub_chunks)
            
            return chunks
        
        # 3. Cümle bazlı bölme (Türkçe cümle sonları)
        # Türkçe için: . ! ? ve ardından boşluk veya satır sonu
        sentence_pattern = r'(?<=[.!?])\s+'
        sentences = re.split(sentence_pattern, text)
        
        if len(sentences) > 1:
            current = ""
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                test_len = len(current) + len(sentence) + 1
                
                if test_len <= max_size:
                    current += (" " if current else "") + sentence
                else:
                    if current and len(current) >= self.MIN_CHUNK_SIZE:
                        chunks.append(current)
                    current = sentence
            
            if current and len(current) >= self.MIN_CHUNK_SIZE:
                chunks.append(current)
            
            return chunks
        
        # 4. Son çare: Karakter bazlı bölme (kelime sınırlarını koruyarak)
        words = text.split()
        current = ""
        
        for word in words:
            test_len = len(current) + len(word) + 1
            
            if test_len <= max_size:
                current += (" " if current else "") + word
            else:
                if current and len(current) >= self.MIN_CHUNK_SIZE:
                    chunks.append(current)
                current = word
        
        if current and len(current) >= self.MIN_CHUNK_SIZE:
            chunks.append(current)
        
        return chunks
    
    def _chunks_from_document(self, doc) -> List[Dict[str, Any]]:
        """
        Document'tan semantik chunk'lar oluşturur.
        
        Strateji:
        1. Başlıklara göre section'lara ayır (heading hiyerarşi korunur)
        2. Her section'ı max size'a göre recursive böl
        3. Tabloları satır bazlı chunk'la (yapısal metadata ile)
        
        v2.43.0: heading_path (breadcrumb), heading_level, tablo metadata.
        """
        chunks = []
        chunk_index = 0
        
        # Heading hiyerarşi koruması (Faz 3)
        heading_stack = []  # [(level, text), ...]
        current_section = []
        current_heading = None
        current_heading_level = 0
        
        def _get_heading_level(style_name: str) -> int:
            """Style name'den heading level çıkarır."""
            if not style_name:
                return 0
            # 'Heading 1' → 1, 'Heading 2' → 2, ...
            for i in range(1, 7):
                if style_name == f'Heading {i}':
                    return i
            if style_name.startswith('Heading'):
                return 1  # Genel heading
            return 0
        
        def _get_heading_path() -> list:
            """Heading stack'ten breadcrumb path oluşturur."""
            return [h[1] for h in heading_stack]
        
        # 🆕 v2.43.0 Faz 4: Paragraf bazlı görsel pozisyonlarını topla
        image_para_indices = []
        for para_idx, para in enumerate(doc.paragraphs):
            # inline_shapes veya XML tabanlı görsel tespiti
            try:
                if para._element is not None:
                    # drawing veya pict elementleri → görsel var
                    ns = {'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                          'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
                    drawings = para._element.findall('.//wp:inline', ns) + para._element.findall('.//wp:anchor', ns)
                    if drawings:
                        image_para_indices.append(para_idx)
            except Exception:
                pass
        
        # Heading-image eşleme: her heading'e ait görselleri belirle
        heading_image_map = {}  # heading_text -> [para_indices with images]
        current_map_heading = None
        for para_idx, para in enumerate(doc.paragraphs):
            is_heading = para.style and para.style.name.startswith('Heading')
            if is_heading and para.text.strip():
                current_map_heading = para.text.strip()
                if current_map_heading not in heading_image_map:
                    heading_image_map[current_map_heading] = []
            elif para_idx in image_para_indices and current_map_heading:
                heading_image_map[current_map_heading].append({
                    "para_index": para_idx,
                    "page": 1,  # DOCX'te sayfa bilgisi yok, yaklaşık
                })
        
        OVERLAP_SIZE = 100  # v3.2.0 RAG-1: Chunk overlap
        
        def _save_section():
            """Mevcut section'ı chunk'lara çevir ve kaydet."""
            nonlocal chunk_index
            if current_section:
                section_text = "\n\n".join(current_section)
                section_text = self._fix_turkish_chars(section_text)
                
                sub_chunks = self._split_large_chunk(section_text)
                for i, sub_text in enumerate(sub_chunks):
                    if len(sub_text.strip()) >= self.MIN_CHUNK_SIZE:
                        # v3.2.0 RAG-1: Overlap — önceki chunk'ın sonundan bağlam ekle
                        overlap_prefix = ""
                        if i > 0 and len(sub_chunks[i - 1]) > OVERLAP_SIZE:
                            overlap_prefix = sub_chunks[i - 1][-OVERLAP_SIZE:].strip() + "\n"
                        
                        final_text = (overlap_prefix + sub_text).strip() if overlap_prefix else sub_text
                        
                        chunk_meta = {
                            "type": "paragraph",
                            "heading": current_heading,
                            "heading_level": current_heading_level,
                            "heading_path": _get_heading_path(),
                            "file_type": "docx",
                            "chunk_index": chunk_index
                        }
                        
                        # 🆕 v2.43.0 Faz 4: Heading-image eşlemesi
                        if current_heading and current_heading in heading_image_map:
                            chunk_meta["image_refs"] = heading_image_map[current_heading]
                        
                        chunks.append({
                            "text": final_text,
                            "metadata": chunk_meta
                        })
                        chunk_index += 1
        
        # Paragrafları işle
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Başlık kontrolü — Style bazlı
            is_heading = para.style and para.style.name.startswith('Heading')
            detected_level = 0
            
            # v3.4.1: Style Heading değilse, font özelliklerine bak
            if not is_heading and len(text) < 100 and not text.endswith('.'):
                try:
                    runs = para.runs
                    if runs:
                        # Tüm run'lar bold mu? (boş run listesinde false-positive önleme)
                        text_runs = [r for r in runs if r.text.strip()]
                        all_bold = bool(text_runs) and all(r.bold for r in text_runs)
                        # İlk run'ın font boyutu
                        first_run = runs[0]
                        font_size = first_run.font.size
                        font_color = first_run.font.color
                        
                        # Renkli text (body text genelde siyah/otomatik)
                        has_color = False
                        if font_color and font_color.rgb:
                            color_val = str(font_color.rgb)
                            has_color = color_val not in ('000000', 'FFFFFF', 'None')
                        
                        # Bold + kısa satır → heading
                        if all_bold and len(text) < 80:
                            is_heading = True
                            detected_level = 3
                        # Renkli + kısa satır → heading
                        elif has_color and len(text) < 80:
                            is_heading = True
                            detected_level = 3
                        # Büyük font (14pt+) + kısa → heading
                        elif font_size and font_size.pt >= 14 and len(text) < 100:
                            is_heading = True
                            detected_level = 2
                except Exception:
                    pass
                
                # Title Case tespiti (fallback)
                if not is_heading and len(text) < 60:
                    words = text.split()
                    if 2 <= len(words) <= 12:
                        tc_count = sum(1 for w in words if len(w) > 1 and w[0].isupper())
                        if tc_count / len(words) >= 0.7:
                            is_heading = True
                            detected_level = 3
            
            if is_heading:
                # Önceki section'ı kaydet
                _save_section()
                
                level = _get_heading_level(para.style.name) if para.style else 0
                if level == 0:
                    level = detected_level or 3  # Font-based tespit
                
                # Heading stack güncelle
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, text))
                
                current_heading = text
                current_heading_level = level
                current_section = []
            else:
                current_section.append(text)
        
        # Son section'ı kaydet
        _save_section()
        
        # Tabloları işle — yapısal metadata ile (Faz 5)
        # En son heading context'ini tabloların heading'i olarak kullan
        last_heading = current_heading
        last_heading_level = current_heading_level
        last_heading_path = _get_heading_path()
        
        for table_idx, table in enumerate(doc.tables):
            if len(table.rows) < 2:
                continue
            
            headers = [cell.text.strip() or f"Sütun{i+1}" for i, cell in enumerate(table.rows[0].cells)]
            row_count = len(table.rows) - 1  # Header hariç
            
            for row_idx, row in enumerate(table.rows[1:], start=2):
                row_values = [cell.text.strip() for cell in row.cells]
                
                if not any(row_values):
                    continue
                
                formatted_lines = []
                for header, value in zip(headers, row_values):
                    if value and value.lower() not in ['none', 'nan', '']:
                        formatted_lines.append(f"**{header}:** {value}")
                
                if formatted_lines:
                    chunk_text = "\n".join(formatted_lines)
                    chunk_text = self._fix_turkish_chars(chunk_text)
                    
                    # Tablo metadata (Faz 5 — yapısal zenginleştirme)
                    table_meta = {
                        "type": "table_row",
                        "table_id": table_idx + 1,
                        "table": table_idx + 1,  # Geriye uyumluluk
                        "row": row_idx,
                        "row_count": row_count,
                        "column_headers": headers,
                        "heading": last_heading,
                        "heading_level": last_heading_level,
                        "heading_path": last_heading_path,
                        "chunk_index": chunk_index
                    }
                    
                    if len(chunk_text) > self.MAX_CHUNK_SIZE:
                        for sub_text in self._split_large_chunk(chunk_text):
                            chunks.append({
                                "text": sub_text,
                                "metadata": {**table_meta, "chunk_index": chunk_index}
                            })
                            chunk_index += 1
                    else:
                        chunks.append({
                            "text": chunk_text,
                            "metadata": table_meta
                        })
                        chunk_index += 1
        
        return chunks
    
    def _extract_from_document(self, doc) -> str:
        """Document objesinden metin çıkarır (legacy - okunabilir format)"""
        chunks = self._chunks_from_document(doc)
        return "\n\n---\n\n".join([c["text"] for c in chunks])
    
    def get_metadata(self, file_path: Path = None, file_name: str = None) -> dict:
        """DOCX metadata'sını çıkarır"""
        base_meta = {"processor": self.PROCESSOR_NAME}
        
        if file_path:
            try:
                from docx import Document
                doc = Document(str(file_path))
                props = doc.core_properties
                
                base_meta.update({
                    "title": props.title or file_path.stem,
                    "author": props.author or "",
                    "created": str(props.created) if props.created else "",
                })
            except Exception:
                logger.warning("[DOCXProcessor] Metadata okuma hatası", exc_info=True)
                base_meta["title"] = file_path.stem
        else:
            base_meta["file_name"] = file_name
        
        return base_meta
